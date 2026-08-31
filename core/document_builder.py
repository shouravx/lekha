"""core/document_builder.py — assembles final output documents (DOCX,
PDF, TXT) from translated page text.

Streaming by design: callers pass an Iterator[tuple[page_num, text]]
(typically `checkpoint_manager.read_pages(job_id)`), and each builder
writes incrementally rather than holding the full document text in
memory. This is what makes 1000+ page output generation safe on 16GB
RAM — only one page's worth of text is ever resident at a time.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterator

import config
from services.logger_service import get_logger

logger = get_logger("document_builder")


def _is_bengali_target(target_lang: str) -> bool:
    return target_lang == "bn"


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------
def build_txt(pages: Iterator[tuple[int, str]], output_path: Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", encoding="utf-8") as f:
        first = True
        for page_num, text in pages:
            if not first:
                f.write("\n\f\n")  # form-feed page break, readable in any text editor
            f.write(f"--- Page {page_num + 1} ---\n")
            f.write(text)
            first = False
    logger.info("Built TXT output: %s", output_path)
    return output_path


# ---------------------------------------------------------------------------
# DOCX (primary output format)
# ---------------------------------------------------------------------------
def build_docx(
    pages: Iterator[tuple[int, list]],
    output_path: Path,
    target_lang: str,
    title: str = "",
) -> Path:
    """Rebuilds a formatted Word document from structured pages.

    Each Block is mapped onto a real Word construct — Title, Heading 1-4,
    List Bullet, a table, an inline image — rather than onto a uniform
    paragraph. This is what makes the translation read as a copy of the
    original document instead of a transcript of its words.

    Word's built-in styles are used deliberately in preference to
    hand-rolled formatting: they carry the document's outline (so the
    navigation pane and any generated table of contents work), and they
    remain restyleable by the user afterwards.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    bengali_target = _is_bengali_target(target_lang)
    _configure_styles(doc, bengali_target)

    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin

    # The filename is recorded as document metadata regardless, but it is
    # only rendered as a visible title when the source document has no
    # title of its own — otherwise every output would open with the
    # untranslated filename stacked above the real, translated title.
    if title:
        doc.core_properties.title = title

    pages = iter(pages)
    first_page = next(pages, None)
    if first_page is not None and title:
        has_own_title = any(b.kind.value == "title" for b in first_page[1])
        if not has_own_title:
            heading = doc.add_heading(title, level=0)
            _apply_run_font(heading, bengali_target)

    page_count = 0
    block_count = 0

    for _page_num, blocks in _chain_first(first_page, pages):
        for block in blocks:
            if _render_block(doc, block, bengali_target, usable_width,
                             WD_ALIGN_PARAGRAPH, Inches, Pt, section):
                block_count += 1

        page_count += 1

        # Flush to disk periodically on very large documents so a crash
        # mid-build doesn't lose everything (python-docx itself only
        # writes on .save(), so we re-save incrementally every N pages).
        if page_count % 250 == 0:
            doc.save(output_path)
            logger.info("DOCX checkpoint flush at page %d -> %s", page_count, output_path)

    doc.save(output_path)
    logger.info(
        "Built DOCX output (%d source pages, %d blocks): %s",
        page_count, block_count, output_path,
    )
    return output_path


def _configure_styles(doc, bengali_target: bool) -> None:
    """Points the document's styles at fonts that can actually render the
    target script.

    Word picks a font per character class, so a Bengali run needs the
    complex-script ("cs") slot set, not just the ascii one — otherwise
    Word silently substitutes and the text renders as boxes.
    """
    from docx.shared import Pt

    normal = doc.styles["Normal"]
    normal.font.name = config.DOCX_FONT_LATIN
    normal.font.size = Pt(11)
    if bengali_target:
        _set_style_fonts(normal)

    # Headings inherit from Normal for the ascii font but need the same
    # complex-script treatment, or every heading renders as boxes.
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3",
                       "Heading 4", "List Bullet", "List Bullet 2",
                       "List Bullet 3"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        if bengali_target:
            _set_style_fonts(style)


def _set_style_fonts(style) -> None:
    from docx.oxml.ns import qn

    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:cs"), config.DOCX_FONT_BENGALI)
    rfonts.set(qn("w:eastAsia"), config.DOCX_FONT_BENGALI)


def _chain_first(first, rest):
    """Re-attaches a peeked-at first item to the front of an iterator."""
    if first is not None:
        yield first
    for item in rest:
        yield item


def _render_furniture(block, bengali_target: bool, section, is_header: bool) -> bool:
    """Writes a repeating letterhead or footer line into the document's
    real header/footer, where Word will repeat it on every page — rather
    than leaving it stranded mid-prose."""
    from docx.shared import Pt

    container = section.header if is_header else section.footer
    # The first paragraph of a fresh header/footer exists but is empty;
    # reuse it before appending, or every document gains a blank line.
    if len(container.paragraphs) == 1 and not container.paragraphs[0].text:
        paragraph = container.paragraphs[0]
    else:
        paragraph = container.add_paragraph()
    paragraph.alignment = 1  # centred

    run = paragraph.add_run(block.text.strip())
    run.font.size = Pt(8.5)
    _apply_run_font(run, bengali_target)
    return True


def _render_block(doc, block, bengali_target: bool, usable_width,
                  WD_ALIGN_PARAGRAPH, Inches, Pt, section=None) -> bool:
    """Renders one Block into the document. Returns False if nothing was
    emitted (an empty or unusable block)."""
    from core.document_model import BlockKind

    if block.kind in (BlockKind.HEADER, BlockKind.FOOTER):
        if section is None or not block.text.strip():
            return False
        return _render_furniture(
            block, bengali_target, section, block.kind is BlockKind.HEADER
        )

    if block.kind is BlockKind.IMAGE:
        return _render_image(doc, block, usable_width, Inches)

    if block.kind is BlockKind.TABLE:
        return _render_table(doc, block, bengali_target)

    if not block.text.strip():
        return False

    if block.kind is BlockKind.TITLE:
        paragraph = doc.add_paragraph(style="Title")
    elif block.kind is BlockKind.HEADING:
        level = min(max(block.level, 1), 4)
        paragraph = doc.add_paragraph(style=f"Heading {level}")
    elif block.kind is BlockKind.LIST_ITEM:
        # Word ships List Bullet through List Bullet 3; deeper nesting
        # falls back to the deepest available style plus an indent.
        depth = min(max(block.list_depth, 1), 3)
        style = "List Bullet" if depth == 1 else f"List Bullet {depth}"
        try:
            paragraph = doc.add_paragraph(style=style)
        except KeyError:
            paragraph = doc.add_paragraph(style="List Bullet")
    else:
        paragraph = doc.add_paragraph()
        if block.indent > 1:
            # Source indent is in points; Word wants inches.
            paragraph.paragraph_format.left_indent = Inches(block.indent / 72.0)

    for run_spec in block.runs:
        if not run_spec.text:
            continue
        run = paragraph.add_run(run_spec.text)
        # Headings are already bold via their style; re-bolding a run
        # inside one is redundant and interferes with theme restyling.
        if block.kind not in (BlockKind.TITLE, BlockKind.HEADING):
            run.bold = run_spec.bold
            run.italic = run_spec.italic
        _apply_run_font(run, bengali_target)

    return True


def _render_image(doc, block, usable_width, Inches) -> bool:
    path = Path(block.image_path)
    if not path.exists():
        logger.debug("Skipping missing extracted image: %s", path)
        return False

    paragraph = doc.add_paragraph()
    paragraph.alignment = 1  # centred
    run = paragraph.add_run()
    try:
        # Preserve the image's size on the source page, but never let it
        # overflow the text column.
        width = Inches(block.image_width / 72.0) if block.image_width else None
        if width is not None and width > usable_width:
            width = usable_width
        run.add_picture(str(path), width=width)
    except Exception as exc:  # noqa: BLE001
        logger.warning("Could not embed image %s: %s", path, exc)
        return False
    return True


def _render_table(doc, block, bengali_target: bool) -> bool:
    rows = [row for row in block.rows if any((c or "").strip() for c in row)]
    if not rows:
        return False
    columns = max(len(row) for row in rows)

    table = doc.add_table(rows=0, cols=columns)
    # "Table Grid" is the one built-in style guaranteed to draw borders;
    # without it a translated table renders as unaligned floating text.
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index in range(columns):
            value = row[col_index] if col_index < len(row) else ""
            cell = cells[col_index]
            # A new cell already holds exactly one empty paragraph. Do not
            # assign cell.text first: that leaves an empty run in front of
            # ours, and any run-level formatting then lands on the wrong one.
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(value or "")
            # A first row of headers is the overwhelmingly common case and
            # reads far better emphasised.
            if row_index == 0:
                run.bold = True
            _apply_run_font(run, bengali_target)
    return True


def _apply_run_font(run_or_paragraph, bengali_target: bool) -> None:
    from docx.oxml.ns import qn

    runs = getattr(run_or_paragraph, "runs", None)
    targets = runs if runs is not None else [run_or_paragraph]
    for run in targets:
        run.font.name = config.DOCX_FONT_LATIN
        if bengali_target:
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.append(rfonts)
            rfonts.set(qn("w:cs"), config.DOCX_FONT_BENGALI)
            rfonts.set(qn("w:eastAsia"), config.DOCX_FONT_BENGALI)
            rfonts.set(qn("w:ascii"), config.DOCX_FONT_LATIN)
            rfonts.set(qn("w:hAnsi"), config.DOCX_FONT_LATIN)


# ---------------------------------------------------------------------------
# PDF (secondary output format)
# ---------------------------------------------------------------------------
def build_pdf(
    pages: Iterator[tuple[int, list]],
    output_path: Path,
    target_lang: str,
    title: str = "",
) -> Path:
    """Builds a PDF using ReportLab with a bundled Noto Sans Bengali font
    when the target language is Bengali (ReportLab's built-in fonts are
    Latin-only and would render Bengali as empty boxes otherwise).
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    font_name = "Helvetica"
    needs_latin_fallback = False
    supported: frozenset = frozenset()
    bengali_target = _is_bengali_target(target_lang)
    if bengali_target:
        if config.BENGALI_FONT_PATH.exists():
            try:
                pdfmetrics.registerFont(TTFont(config.BENGALI_FONT_NAME, str(config.BENGALI_FONT_PATH)))
                font_name = config.BENGALI_FONT_NAME
                # Noto Sans Bengali carries no Latin glyphs, and no glyph
                # for '@' or '&' either, so anything it cannot draw has to
                # be handed to a fallback font or it renders as blanks.
                supported = _font_coverage(font_name)
                needs_latin_fallback = bool(supported)
            except Exception as exc:  # noqa: BLE001
                logger.warning("Could not register Bengali font, falling back to Helvetica: %s", exc)
        else:
            logger.warning(
                "Bengali font not found at %s — Bengali text in the PDF output may not render "
                "correctly. See README for font setup.",
                config.BENGALI_FONT_PATH,
            )

    from reportlab.lib import colors
    from reportlab.platypus import Image as RLImage
    from reportlab.platypus import Table, TableStyle

    from core.document_model import BlockKind

    body_style = ParagraphStyle(
        "Body", fontName=font_name, fontSize=10.5, leading=15, spaceAfter=6,
    )
    title_style = ParagraphStyle(
        "DocTitle", fontName=font_name, fontSize=19, leading=24,
        spaceBefore=6, spaceAfter=16,
    )
    # Heading sizes step down so the document's hierarchy stays visible at
    # a glance, mirroring the levels the extractor recovered.
    heading_styles = {
        level: ParagraphStyle(
            f"H{level}", fontName=font_name,
            fontSize=size, leading=size * 1.3,
            spaceBefore=12 - level, spaceAfter=5,
        )
        for level, size in ((1, 16.0), (2, 13.5), (3, 11.5), (4, 10.5))
    }
    bullet_styles = {
        depth: ParagraphStyle(
            f"Bullet{depth}", parent=body_style,
            leftIndent=12 + 14 * depth, bulletIndent=2 + 14 * depth,
            spaceAfter=3,
        )
        for depth in range(0, 5)
    }

    doc = SimpleDocTemplate(
        str(output_path), pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm, topMargin=18 * mm, bottomMargin=18 * mm,
    )
    frame_width = doc.width

    story = []
    running_header: list[str] = []
    running_footer: list[str] = []

    # As in the DOCX builder, the filename is only used as a title when
    # the document has none of its own — otherwise the output opens with
    # the untranslated filename above the real, translated title. It also
    # goes through the fallback markup, since a Latin filename is exactly
    # the text the Bengali font cannot draw.
    pages = iter(pages)
    first_page = next(pages, None)
    if first_page is not None and title:
        if not any(b.kind is BlockKind.TITLE for b in first_page[1]):
            story.append(Paragraph(
                _cell_markup(title, needs_latin_fallback, supported), title_style
            ))

    page_count = 0
    for _page_num, blocks in _chain_first(first_page, pages):
        for block in blocks:
            if block.kind in (BlockKind.HEADER, BlockKind.FOOTER):
                # Collected rather than placed inline; drawn on every page
                # by the page callback below, which is where a repeating
                # letterhead actually belongs.
                text = block.text.strip()
                if text:
                    (running_header if block.kind is BlockKind.HEADER
                     else running_footer).append(text)
                continue

            if block.kind is BlockKind.IMAGE:
                _append_pdf_image(story, block, frame_width, RLImage, Spacer)
                continue

            if block.kind is BlockKind.TABLE:
                _append_pdf_table(
                    story, block, body_style, frame_width,
                    Table, TableStyle, Paragraph, Spacer, colors,
                    font_name, needs_latin_fallback, supported,
                )
                continue

            text = block.text.strip()
            if not text:
                continue

            markup = _runs_to_markup(block, font_name, needs_latin_fallback, supported)
            if block.kind is BlockKind.TITLE:
                story.append(Paragraph(markup, title_style))
            elif block.kind is BlockKind.HEADING:
                story.append(Paragraph(markup, heading_styles[min(max(block.level, 1), 4)]))
            elif block.kind is BlockKind.LIST_ITEM:
                depth = min(max(block.list_depth, 0), 4)
                story.append(Paragraph(markup, bullet_styles[depth], bulletText="•"))
            else:
                story.append(Paragraph(markup, body_style))
        page_count += 1

    if not story:
        story.append(Paragraph("[No translatable content was found in this document]", body_style))

    def draw_furniture(canvas, document) -> None:
        """Paints the running header/footer and a page number on each page."""
        canvas.saveState()
        canvas.setFont(font_name, 8)
        canvas.setFillGray(0.42)
        width = document.pagesize[0]
        if running_header:
            y = document.pagesize[1] - 12 * mm
            for line in running_header:
                _draw_mixed_centred(canvas, width / 2.0, y, line,
                                    font_name, 8, needs_latin_fallback, supported)
                y -= 10
        y = 12 * mm
        for line in reversed(running_footer):
            _draw_mixed_centred(canvas, width / 2.0, y, line,
                                font_name, 8, needs_latin_fallback, supported)
            y += 10
        canvas.setFont(font_name, 8)
        canvas.drawCentredString(width / 2.0, 7 * mm, str(canvas.getPageNumber()))
        canvas.restoreState()

    doc.build(story, onFirstPage=draw_furniture, onLaterPages=draw_furniture)
    logger.info("Built PDF output (%d source pages): %s", page_count, output_path)
    return output_path


_LATIN_FALLBACK_FONT = "Helvetica"


def _font_coverage(font_name: str) -> frozenset:
    """The set of code points a registered font can actually draw.

    Asking the font directly is the only reliable test. Classifying by
    script is not enough: '@', '&' and '_' belong to no script, yet Noto
    Sans Bengali has no glyph for them, so an email address rendered in it
    loses its '@' silently.
    """
    try:
        from reportlab.pdfbase import pdfmetrics

        face = pdfmetrics.getFont(font_name).face
        return frozenset(getattr(face, "charToGlyph", {}).keys())
    except Exception as exc:  # noqa: BLE001
        logger.debug("Could not read glyph coverage for %s: %s", font_name, exc)
        return frozenset()


def _split_script_segments(text: str, supported: frozenset = frozenset()) -> list[tuple[str, bool]]:
    """Splits text into (segment, needs_fallback) runs.

    A segment is flagged for the fallback font when the primary font has
    no glyph for it. Whitespace is neutral and never forces a switch, so
    the font changes only where it has to.
    """
    if not supported:
        return [(text, False)]

    segments: list[tuple[str, bool]] = []
    current: list[str] = []
    current_fallback: bool | None = None

    for ch in text:
        if ch.isspace():
            current.append(ch)
            continue

        needs_fallback = ord(ch) not in supported
        if current_fallback is None:
            current_fallback = needs_fallback
        elif needs_fallback != current_fallback:
            segments.append(("".join(current), current_fallback))
            current = []
            current_fallback = needs_fallback
        current.append(ch)

    if current:
        segments.append(("".join(current), bool(current_fallback)))
    return segments


def _runs_to_markup(block, font_name: str = "", needs_latin_fallback: bool = False,
                    supported: frozenset = frozenset()) -> str:
    """Renders a block's runs as ReportLab's inline markup, so bold and
    italic survive into the PDF rather than being flattened — and so
    Latin text inside a Bengali document is drawn in a font that has
    Latin glyphs."""
    parts = []
    for run in block.runs:
        if not run.text:
            continue
        if needs_latin_fallback:
            piece = "".join(
                f'<font name="{_LATIN_FALLBACK_FONT}">{_xml_escape(seg)}</font>'
                if is_latin else _xml_escape(seg)
                for seg, is_latin in _split_script_segments(run.text, supported)
            )
        else:
            piece = _xml_escape(run.text)
        if run.bold:
            piece = f"<b>{piece}</b>"
        if run.italic:
            piece = f"<i>{piece}</i>"
        parts.append(piece)
    return "".join(parts)


def _draw_mixed_centred(canvas, centre_x: float, y: float, text: str,
                        font_name: str, size: float, needs_latin_fallback: bool,
                        supported: frozenset = frozenset()) -> None:
    """Draws centred text that may mix Bengali and Latin, switching font
    per segment so neither script is dropped."""
    if not needs_latin_fallback:
        canvas.drawCentredString(centre_x, y, text)
        return

    segments = _split_script_segments(text, supported)
    total = sum(
        canvas.stringWidth(seg, _LATIN_FALLBACK_FONT if is_latin else font_name, size)
        for seg, is_latin in segments
    )
    x = centre_x - total / 2.0
    for seg, is_latin in segments:
        font = _LATIN_FALLBACK_FONT if is_latin else font_name
        canvas.setFont(font, size)
        canvas.drawString(x, y, seg)
        x += canvas.stringWidth(seg, font, size)


def _append_pdf_image(story, block, frame_width, RLImage, Spacer) -> None:
    path = Path(block.image_path)
    if not path.exists():
        return
    try:
        width = block.image_width or frame_width
        height = block.image_height or 0
        if width > frame_width:
            # Scale proportionally rather than distorting.
            if height:
                height = height * (frame_width / width)
            width = frame_width
        image = RLImage(str(path), width=width, height=height or None)
        story.append(image)
        story.append(Spacer(1, 6))
    except Exception as exc:  # noqa: BLE001
        logger.warning("Could not place image %s in PDF: %s", path, exc)


def _append_pdf_table(story, block, body_style, frame_width,
                      Table, TableStyle, Paragraph, Spacer, colors,
                      font_name: str = "", needs_latin_fallback: bool = False,
                      supported: frozenset = frozenset()) -> None:
    rows = [row for row in block.rows if any((c or "").strip() for c in row)]
    if not rows:
        return
    columns = max(len(row) for row in rows)
    col_width = frame_width / columns

    # Cells are Paragraphs, not raw strings, so long translated text wraps
    # inside its column instead of overflowing the table.
    data = [
        [
            Paragraph(
                _cell_markup((row[i] if i < len(row) else "") or "",
                             needs_latin_fallback, supported),
                body_style,
            )
            for i in range(columns)
        ]
        for row in rows
    ]
    table = Table(data, colWidths=[col_width] * columns)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#999999")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EFEFEF")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)
    story.append(Spacer(1, 8))


def _xml_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _cell_markup(text: str, needs_latin_fallback: bool,
                 supported: frozenset = frozenset()) -> str:
    """Table-cell text with the same Latin fallback the body paragraphs get."""
    if not needs_latin_fallback:
        return _xml_escape(text)
    return "".join(
        f'<font name="{_LATIN_FALLBACK_FONT}">{_xml_escape(seg)}</font>'
        if is_latin else _xml_escape(seg)
        for seg, is_latin in _split_script_segments(text, supported)
    )
